"""Render a grouped report into a .docx with one table per category."""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt


def write_docx(
    path: Path,
    *,
    location_name: str,
    grouped_categories: Iterable[dict],
) -> None:
    doc = Document()
    title = doc.add_heading(f"{location_name} — InnoINVest data report", level=1)
    title.style.font.size = Pt(18)

    for cat in grouped_categories:
        doc.add_heading(cat["category"], level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid"
        header = table.rows[0].cells
        header[0].text = "KPI"
        header[1].text = "Value"
        header[2].text = "Year"
        header[3].text = "Source"
        for kpi in cat["kpis"]:
            row = table.add_row().cells
            unit = kpi["unit"]
            latest = kpi["latest"]
            row[0].text = kpi.get("kpi_name_en") or kpi.get("name_en", "")
            row[1].text = _format_value(latest["value"], unit)
            row[2].text = latest["period"]
            row[3].text = _format_source(latest)
        doc.add_paragraph()

    doc.save(str(path))


def _format_value(value: str | None, unit: str) -> str:
    if value is None:
        return "—"
    suffix = {"EUR": " €", "RON": " RON", "percent": " %", "persons": ""}.get(unit, f" {unit}")
    return f"{value}{suffix}"


def _format_source(latest: dict) -> str:
    src = latest["source_code"].upper()
    if latest.get("source_dataset_id"):
        return f"{src} {latest['source_dataset_id']}"
    return src
