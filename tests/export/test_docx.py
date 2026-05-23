from pathlib import Path

from docx import Document

from innoinvest.export.docx import write_docx


def test_write_docx_produces_grouped_tables(tmp_path: Path):
    out = tmp_path / "out.docx"
    write_docx(out, location_name="Floresti", grouped_categories=[
        {"category": "Demographics", "kpis": [
            {"kpi_name_en": "Total population", "unit": "persons",
             "latest": {"period": "2025", "value": "57437",
                        "source_code": "ins_tempo", "source_dataset_id": "POP107D"}},
        ]},
    ])

    assert out.exists() and out.stat().st_size > 0
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Floresti" in text
    assert "Demographics" in text
    table_text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Total population" in table_text
    assert "57437" in table_text
    assert "POP107D" in table_text
